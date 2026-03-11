from __future__ import annotations

import re
from pathlib import Path

from docx import Document

BOLD_PATTERN = re.compile(r"(\*\*.*?\*\*)")
ORDERED_LIST_PATTERN = re.compile(r"^\d+\.\s+")


def _add_inline_runs(paragraph, text: str) -> None:
    parts = BOLD_PATTERN.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def markdown_to_docx(markdown_text: str, output_path: Path) -> Path:
    document = Document()
    lines = markdown_text.splitlines()
    buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal buffer
        if not buffer:
            return
        paragraph = document.add_paragraph()
        _add_inline_runs(paragraph, " ".join(item.strip() for item in buffer).strip())
        buffer = []

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            continue

        if stripped == "---":
            flush_paragraph()
            document.add_paragraph()
            continue

        if stripped.startswith("### "):
            flush_paragraph()
            document.add_heading(stripped[4:].strip(), level=3)
            continue

        if stripped.startswith("## "):
            flush_paragraph()
            document.add_heading(stripped[3:].strip(), level=2)
            continue

        if stripped.startswith("# "):
            flush_paragraph()
            document.add_heading(stripped[2:].strip(), level=1)
            continue

        if stripped.startswith("- "):
            flush_paragraph()
            paragraph = document.add_paragraph(style="List Bullet")
            _add_inline_runs(paragraph, stripped[2:].strip())
            continue

        if ORDERED_LIST_PATTERN.match(stripped):
            flush_paragraph()
            content = ORDERED_LIST_PATTERN.sub("", stripped, count=1)
            paragraph = document.add_paragraph(style="List Number")
            _add_inline_runs(paragraph, content)
            continue

        buffer.append(stripped)

    flush_paragraph()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def export_markdown_file_to_docx(markdown_path: Path, output_path: Path | None = None) -> Path:
    target_path = output_path or markdown_path.with_suffix(".docx")
    markdown_text = markdown_path.read_text(encoding="utf-8")
    return markdown_to_docx(markdown_text, target_path)
